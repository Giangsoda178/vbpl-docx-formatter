"""Generate a sample Vietnamese legal docx for testing the formatter."""
from docx import Document

doc = Document()

# Header table (quoc hieu)
t = doc.add_table(rows=2, cols=2)
t.cell(0, 0).text = 'CHÍNH PHỦ\n--------'
t.cell(0, 1).text = 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n---------------'
t.cell(1, 0).text = 'Số: 141/2026/NĐ-CP'
t.cell(1, 1).text = 'Hà Nội, ngày 29 tháng 4 năm 2026'

doc.add_paragraph('NGHỊ ĐỊNH')
doc.add_paragraph('QUY ĐỊNH VỀ QUẢN LÝ DỮ LIỆU SỐ')
doc.add_paragraph('Căn cứ Luật Tổ chức Chính phủ ngày 19 tháng 6 năm 2015;')
doc.add_paragraph('Theo đề nghị của Bộ trưởng Bộ Thông tin và Truyền thông;')
doc.add_paragraph('Chính phủ ban hành Nghị định quy định về quản lý dữ liệu số.')
doc.add_paragraph('Chương I')
doc.add_paragraph('Quy định chung')
doc.add_paragraph('Điều 1. Phạm vi điều chỉnh')
doc.add_paragraph('Nghị định này quy định về quản lý, kết nối và chia sẻ dữ liệu số.')
doc.add_paragraph('Điều 2. Đối tượng áp dụng')
doc.add_paragraph('1. Cơ quan nhà nước các cấp.')
doc.add_paragraph('2. Tổ chức, cá nhân có liên quan.')

# Closing table
t2 = doc.add_table(rows=1, cols=2)
t2.cell(0, 0).text = 'Nơi nhận:\n- Ban Bí thư Trung ương Đảng;\n- Thủ tướng, các Phó Thủ tướng Chính phủ;\n- Lưu: VT.'
t2.cell(0, 1).text = 'TM. CHÍNH PHỦ\nTHỦ TƯỚNG\nPhạm Minh Chính'

doc.save('test/sample_nghi_dinh.docx')
print('Sample saved.')
