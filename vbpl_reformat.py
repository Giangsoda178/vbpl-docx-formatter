#!/usr/bin/env python3
"""
vbpl_reformat.py
================
Định dạng lại file Word (.docx) văn bản pháp luật Việt Nam đã có sẵn theo
chuẩn định dạng văn bản hành chính (Times New Roman 14pt, justify, thụt
dòng đầu 0.5", các tiêu đề Điều/Chương/Mục đậm, Căn cứ in nghiêng, v.v.).

Hỗ trợ các loại văn bản: Nghị định, Luật, Thông tư, Quyết định, Nghị quyết,
Chỉ thị.

CÁCH DÙNG
---------
    python vbpl_reformat.py <duong_dan_input.docx>
    python vbpl_reformat.py <duong_dan_input.docx> <duong_dan_output.docx>

Nếu không chỉ định output, file kết quả sẽ là <input>_formatted.docx
cùng thư mục với input.

VÍ DỤ
-----
    python vbpl_reformat.py "Nghị định 141.docx"
    → Xuất "Nghị định 141_formatted.docx"

YÊU CẦU
-------
    pip install python-docx

CƠ CHẾ
------
Script đọc file đầu vào, trích xuất:
  1. Metadata từ bảng quốc hiệu (cơ quan, số văn bản, ngày)
  2. Loại văn bản + tiêu đề
  3. Danh sách các đoạn "Căn cứ"
  4. Các Chương / Mục / Điều / khoản trong thân bài
  5. Khối Nơi nhận + chữ ký

Sau đó xây dựng lại văn bản mới với định dạng chuẩn.
"""

import argparse
import copy
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import (WD_CELL_VERTICAL_ALIGNMENT,
                                  WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.exceptions import PackageNotFoundError
    from docx.text.paragraph import Paragraph
except ImportError as exc:
    raise ImportError(
        "The python-docx package is required. Install it with: pip install python-docx"
    ) from exc


# ============================================================
# HẰNG SỐ ĐỊNH DẠNG — chỉnh ở đây để đổi style toàn văn bản
# ============================================================

DEFAULT_FONT = "Times New Roman"
SIZE_DEFAULT = 12       # Quốc hiệu, NGHỊ ĐỊNH/LUẬT, chữ ký
SIZE_BODY = 14          # Thân bài, Điều, Căn cứ, tiêu đề dài
SIZE_HEADING = 12       # Trích yếu, Chương, Mục, mục La Mã, PHỤ LỤC (Điều giữ 14)
SIZE_DOC_NUMBER = 13    # Số văn bản
SIZE_RECIPIENTS = 11    # Danh sách Nơi nhận
SIZE_TABLE = 12         # Nội dung bảng dữ liệu (Phụ lục)
SIZE_MAU = SIZE_BODY    # Nội dung Mẫu biểu — giữ nguyên định dạng gốc,
                        # chỉ chuẩn hóa cỡ chữ (đổi giá trị này để đổi cỡ chữ Mẫu)
TABLE_ROW_TOP_PAD = 120  # twip (~6pt) — đệm phía trên mỗi hàng bảng cho thoáng

INDENT_FIRST = Cm(1.27)   # Thụt dòng đầu = 0.5"
SPACE_BEFORE = Pt(3)      # Khoảng cách trên mỗi đoạn body
SPACE_ZERO = Pt(0)

PAGE_WIDTH = Cm(21.59)    # 8.5"
PAGE_HEIGHT = Cm(27.94)   # 11"
MARGIN_TOP = Cm(1.5)      # Lề trên (khớp mẫu desired ND_so_132)
MARGIN_BOTTOM = Cm(1.5)   # Lề dưới
MARGIN_OTHER = Cm(2.54)   # Lề trái/phải = 1"

# Các loại văn bản nhận dạng được
LOAI_VAN_BAN_KEYWORDS = {
    'NGHỊ ĐỊNH', 'LUẬT', 'THÔNG TƯ', 'QUYẾT ĐỊNH',
    'CHỈ THỊ', 'NGHỊ QUYẾT', 'PHÁP LỆNH',
    'THÔNG TƯ LIÊN TỊCH',
}

# Tên cơ quan ban hành đứng riêng 1 dòng (đậm, căn giữa) — vd "QUỐC HỘI",
# "CHÍNH PHỦ". Nhận dạng theo danh sách + đối chiếu với cơ quan ở quốc hiệu.
CO_QUAN_BODY_KEYWORDS = {
    'QUỐC HỘI', 'CHÍNH PHỦ', 'ỦY BAN THƯỜNG VỤ QUỐC HỘI',
    'THỦ TƯỚNG CHÍNH PHỦ', 'CHỦ TỊCH NƯỚC',
}

# Câu lệnh ban hành đứng riêng (đậm, căn giữa): "QUYẾT NGHỊ:", "NGHỊ ĐỊNH:",
# "QUYẾT ĐỊNH:", "BAN HÀNH:" — thường nằm ngay trước phần thân điều khoản.
ENACTING_KEYWORDS = {
    'QUYẾT NGHỊ:', 'QUYẾT ĐỊNH:', 'NGHỊ ĐỊNH:', 'BAN HÀNH:',
    'QUYẾT NGHỊ', 'QUYẾT ĐỊNH', 'NGHỊ ĐỊNH', 'BAN HÀNH',
}

# Tiêu đề loại văn bản đính kèm đứng riêng 1 dòng (đậm, căn giữa) — phần văn
# bản được "ban hành kèm theo" một Quyết định: vd "KẾ HOẠCH", "CHƯƠNG TRÌNH",
# "ĐỀ ÁN". Xử lý GIỐNG "PHỤ LỤC": sang trang mới, tiêu đề đậm căn giữa, dòng
# kế tiếp là tiêu đề phụ + dòng "(Kèm theo Quyết định số ...)".
ATTACHED_DOC_KEYWORDS = {
    'KẾ HOẠCH', 'CHƯƠNG TRÌNH', 'ĐỀ ÁN', 'QUY CHẾ', 'QUY ĐỊNH',
    'PHƯƠNG ÁN', 'ĐIỀU LỆ', 'CHIẾN LƯỢC', 'QUY HOẠCH',
}

# Tiêu đề Mẫu biểu trong Phụ lục đứng riêng 1 dòng (vd "Mẫu 01. QĐKTĐL",
# "Mẫu số 02. BBKT"). Từ dòng này trở đi (đến hết văn bản), toàn bộ các Mẫu
# cùng nội dung của chúng được GIỮ NGUYÊN định dạng gốc — chỉ chuẩn hóa font
# chữ (Times New Roman) và cỡ chữ — thay vì dựng lại theo style chuẩn.
MAU_HEADING_RE = re.compile(r'^\s*Mẫu\s+(số\s+)?\d+', re.IGNORECASE)

# Tiêu đề danh mục mẫu biểu đứng riêng (vd "DANH MỤC 1", "DANH MỤC 2") —
# mỗi danh mục bắt đầu một trang mới (giống Phụ lục / Mẫu số). Yêu cầu có số
# theo sau để không nhầm với tiêu đề con kiểu "DANH MỤC XUẤT BẢN PHẨM...".
DANH_MUC_HEADING_RE = re.compile(r'^DANH MỤC\s+\d', re.IGNORECASE)

# Các đoạn mở đầu (sẽ in nghiêng + thụt dòng)
CAN_CU_STARTERS = (
    'Căn cứ',
    'Theo đề nghị',
    'Xét đề nghị',
    'Trên cơ sở',
    'Chính phủ ban hành',
    'Quốc hội ban hành',
    'Bộ trưởng ban hành',
    'Thủ tướng Chính phủ ban hành',
)


# ============================================================
# HELPER — Xử lý XML mức thấp
# ============================================================

def _set_cell_borders_none(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    _AFTER = {qn('w:shd'), qn('w:noWrap'), qn('w:tcMar'),
              qn('w:textDirection'), qn('w:tcFitText'),
              qn('w:vAlign'), qn('w:hideMark')}
    insert_at = len(tcPr)
    for i, child in enumerate(tcPr):
        if child.tag in _AFTER:
            insert_at = i
            break
    tcPr.insert(insert_at, tcBorders)


def _set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    _AFTER = {qn('w:shd'), qn('w:tblLayout'), qn('w:tblCellMar'),
              qn('w:tblLook'), qn('w:tblCaption'), qn('w:tblDescription')}
    insert_at = len(tblPr)
    for i, child in enumerate(tblPr):
        if child.tag in _AFTER:
            insert_at = i
            break
    tblPr.insert(insert_at, tblBorders)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders_none(cell)


def _set_table_borders_single(table):
    """Kẻ viền đơn (single line) cho bảng dữ liệu — dùng cho bảng trong
    Phụ lục (khác với bảng quốc hiệu/khối ký vốn không viền)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Xóa viền cũ nếu có
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')        # 0.5pt
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    _AFTER = {qn('w:shd'), qn('w:tblLayout'), qn('w:tblCellMar'),
              qn('w:tblLook'), qn('w:tblCaption'), qn('w:tblDescription')}
    insert_at = len(tblPr)
    for i, child in enumerate(tblPr):
        if child.tag in _AFTER:
            insert_at = i
            break
    tblPr.insert(insert_at, tblBorders)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE \\* MERGEFORMAT'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _set_title_page(section):
    sectPr = section._sectPr
    if sectPr.find(qn('w:titlePg')) is None:
        titlePg = OxmlElement('w:titlePg')
        _AFTER = {qn('w:textDirection'), qn('w:bidi'), qn('w:rtlGutter'),
                  qn('w:docGrid'), qn('w:printerSettings'),
                  qn('w:sectPrChange')}
        insert_at = len(sectPr)
        for i, child in enumerate(sectPr):
            if child.tag in _AFTER:
                insert_at = i
                break
        sectPr.insert(insert_at, titlePg)


def _set_default_font(doc, font_name, size_pt):
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), font_name)


def _add_run(paragraph, text, *, bold=False, italic=False, size=SIZE_BODY):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    return run


def _set_cant_split(table_row):
    """Không cho row bị tách qua 2 trang."""
    tr = table_row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trPr.append(OxmlElement('w:cantSplit'))


# ============================================================
# TRÍCH XUẤT — Đọc dữ liệu có cấu trúc từ document gốc
# ============================================================

def _cell_lines(cell):
    """Lấy các dòng không trống trong cell. Tách cả theo paragraph và
    theo line break (<w:br/>) — vì nhiều file gốc dùng br thay vì
    paragraph riêng."""
    lines = []
    for p in cell.paragraphs:
        for line in p.text.split('\n'):
            line = line.strip()
            if line:
                lines.append(line)
    return lines


def extract_metadata(doc):
    """Trích cơ quan, số văn bản, địa điểm, ngày từ bảng quốc hiệu."""
    if not doc.tables:
        raise ValueError(
            "Không tìm thấy bảng quốc hiệu ở đầu văn bản. "
            "File có vẻ không phải văn bản pháp luật chuẩn.")

    table = doc.tables[0]
    if len(table.rows) < 2:
        raise ValueError("Bảng quốc hiệu cần có ít nhất 2 hàng.")

    # Hàng 1, ô trái: tên cơ quan (có thể kèm dòng "--------"). Tên cơ quan
    # có thể trải trên nhiều dòng (vd "BỘ VĂN HÓA, THỂ THAO" + "VÀ DU LỊCH")
    # → ghép tất cả dòng không phải gạch ngang lại thành tên đầy đủ.
    co_quan_lines = _cell_lines(table.rows[0].cells[0])
    co_quan = ' '.join(l for l in co_quan_lines if not re.fullmatch(r'-+', l))

    # Hàng 2, ô trái: "Số: 141/2026/NĐ-CP" hoặc "Luật số: 58/2024/QH15"
    so_text = ' '.join(_cell_lines(table.rows[1].cells[0]))
    m = re.match(r'^(.+?:)\s*(.+)$', so_text)
    if m:
        prefix_so = m.group(1) + ' '
        so_van_ban = m.group(2)
    else:
        prefix_so = 'Số: '
        so_van_ban = so_text

    # Hàng 2, ô phải: "Hà Nội, ngày 29 tháng 4 năm 2026"
    dia_ngay = ' '.join(_cell_lines(table.rows[1].cells[1]))
    m = re.match(r'^(.+?),\s*ngày\s+(.+)$', dia_ngay)
    if m:
        dia_diem = m.group(1).strip()
        ngay = m.group(2).strip()
    else:
        dia_diem = 'Hà Nội'
        ngay = dia_ngay.strip()

    return {
        'co_quan': co_quan,
        'so_van_ban': so_van_ban,
        'dia_diem': dia_diem,
        'ngay_ban_hanh': ngay,
        'prefix_so': prefix_so,
    }


def _is_para_italic(para):
    """Đoạn có ≥ nửa số run là italic không?"""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    italic_runs = [r for r in runs if r.italic]
    return len(italic_runs) >= len(runs) / 2


def _runs_with_format(para):
    """Lấy list (text, {'bold', 'italic'}) cho từng run của paragraph."""
    parts = []
    for run in para.runs:
        if run.text:
            parts.append((run.text, {
                'bold': bool(run.bold),
                'italic': bool(run.italic),
            }))
    return parts


def _has_mixed_inline_format(para):
    """Đoạn có pha trộn run italic và non-italic không?
    (vd: định nghĩa thuật ngữ '1. *Báo cáo...* là tài liệu...')."""
    runs = _runs_with_format(para)
    if len(runs) < 2:
        return False
    has_i = any(f['italic'] for _, f in runs)
    has_ni = any(not f['italic'] for _, f in runs)
    return has_i and has_ni


def _is_para_bold(para):
    """Đoạn có toàn bộ (≥ nửa) run là bold không? — dùng cho tiêu đề mục
    đậm kiểu '1. Quán triệt...' trong nghị quyết Chính phủ."""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False
    bold_runs = [r for r in runs if r.bold]
    return len(bold_runs) >= len(runs) / 2


def _classify(text, co_quan=''):
    """Phân loại 1 đoạn dựa trên nội dung text.

    co_quan: tên cơ quan lấy từ quốc hiệu — dùng để nhận dạng dòng cơ quan
    ban hành đứng riêng trong thân bài (vd "QUỐC HỘI", "CHÍNH PHỦ")."""
    text = text.strip()
    if not text:
        return ('blank',)

    if text in LOAI_VAN_BAN_KEYWORDS:
        return ('loai_van_ban', text)

    # Dòng câu lệnh ban hành: "QUYẾT NGHỊ:", "NGHỊ ĐỊNH:" ...
    if text in ENACTING_KEYWORDS:
        return ('quyet_nghi', text)

    # Dòng tên cơ quan ban hành đứng riêng (đậm, căn giữa)
    if text in CO_QUAN_BODY_KEYWORDS or (
            co_quan and text.upper() == co_quan.strip().upper()):
        return ('co_quan_body', text)

    # Dòng đơn vị tính của bảng: "Đơn vị: Tỷ đồng", "Đơn vị tính: %"
    if re.match(r'^Đơn vị\s*(tính)?\s*:', text):
        return ('don_vi', text)

    # Dòng tiêu đề mẫu biểu đứng riêng: "Mẫu số 01", "Mẫu số 02" (trong Phụ lục)
    if re.fullmatch(r'Mẫu số\s+\S+', text):
        return ('mau_so', text)

    # Chương / Mục / Điều — chấp nhận dấu mở ngoặc kép đầu dòng (văn bản sửa
    # đổi, bổ sung trích nguyên văn điều/chương được chèn vào luật gốc trong
    # cặp ngoặc kép) và hậu tố chữ cái sau số (vd "Chương Va", "Điều 39a",
    # "Điều 3b", "Điều 39đ"). Nhóm 1 = dấu ngoặc kép mở (nếu có), giữ lại
    # khi dựng để bảo toàn ý nghĩa pháp lý của phần trích dẫn.
    m = re.fullmatch(r'(["“”]?)\s*Chương\s+([IVXLCDM]+[a-zđ]?)', text)
    if m:
        return ('chuong_num', m.group(2), m.group(1))

    m = re.match(r'^(["“”]?)\s*Mục\s+(\d+[a-zđ]?)\.\s+(.+)$', text)
    if m:
        return ('muc', m.group(2), m.group(3).strip(), m.group(1))

    m = re.match(r'^(["“”]?)\s*Điều\s+(\d+[a-zđ]?)\.\s+(.+)$', text)
    if m:
        return ('dieu', m.group(2), m.group(3).strip(), m.group(1))

    # Tiêu đề mục La Mã kiểu Chính phủ: "I. Về quan điểm...", "II. ..."
    # (số La Mã + dấu chấm + tiêu đề trên cùng dòng). Phân biệt với
    # "I)" hay danh sách thường bằng cách yêu cầu dấu chấm + khoảng trắng.
    m = re.match(r'^([IVXLCDM]+)\.\s+(.+)$', text)
    if m:
        return ('roman_section', m.group(1), m.group(2).strip())

    for starter in CAN_CU_STARTERS:
        if text.startswith(starter):
            return ('can_cu', text)

    return ('khoan', text)


def _split_phu_luc_title(para):
    """Tách paragraph tiêu đề phụ lục thành (tiêu_đề, dòng_kèm_theo).

    Tiêu đề phụ lục thường gộp 2 phần qua line break:
        "KẾ HOẠCH VỐN ...\n(Kèm theo Nghị quyết số ... của Quốc hội)"
    Trả về (title, ref) — ref có thể rỗng nếu không có dòng "(Kèm theo...)"."""
    lines = [l.strip() for l in para.text.split('\n') if l.strip()]
    title_lines, ref_lines = [], []
    for l in lines:
        if l.startswith('(') or l.startswith('Kèm theo') or ref_lines:
            ref_lines.append(l)
        else:
            title_lines.append(l)
    return ' '.join(title_lines), ' '.join(ref_lines)


def _is_header_table(table):
    """Bảng có phải khối quốc hiệu (cơ quan + 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA')
    không? — dùng để nhận dạng header của Mẫu biểu trong Phụ lục (render
    không viền) thay vì coi là bảng dữ liệu (có viền)."""
    txt = ' '.join(c.text for row in table.rows for c in row.cells).upper()
    return 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA' in txt


# Dấu hiệu khối chữ ký / Nơi nhận bên trong một Mẫu biểu (Phụ lục).
SIGN_BLOCK_HINTS = (
    'Nơi nhận', 'NGƯỜI ĐẠI DIỆN', 'THỦ TRƯỞNG', 'Người lập báo cáo',
    'Ký và ghi rõ', 'Chữ ký của', 'NGƯỜI ĐỨNG ĐẦU', 'GIÁM ĐỐC',
    'CHỦ TỊCH', 'XÁC NHẬN CỦA',
)


def _is_sign_block(table):
    """Bảng có phải khối chữ ký / Nơi nhận bên trong một Mẫu biểu không?

    Đặc trưng: CHỈ 1 hàng, ≥ 2 ô đặt cạnh nhau (Nơi nhận | chức danh ký),
    chứa dấu hiệu chức danh ký / Nơi nhận. Điều kiện '1 hàng' phân biệt với
    bảng dữ liệu thật (nhiều hàng) — kể cả bảng có cột tiêu đề chứa
    '(Chữ ký...)' như bảng gia hạn giấy phép (≥ 2 hàng → vẫn là bảng dữ liệu).
    """
    if len(table.rows) != 1 or len(table.columns) < 2:
        return False
    txt = ' '.join(c.text for c in table.rows[0].cells)
    return any(h in txt for h in SIGN_BLOCK_HINTS)


def _classify_flow(flow_items, co_quan='', appendix=False):
    """Phân loại 1 chuỗi flow item (paragraph + table) thành list item
    có cấu trúc, giữ nguyên thứ tự xuất hiện của bảng dữ liệu.

    flow_items: list ('p', Paragraph) | ('t', Table).
    co_quan   : tên cơ quan ở quốc hiệu (để nhận dòng cơ quan ban hành).
    appendix  : True khi đang xử lý phần Phụ lục (nhận dạng 'PHỤ LỤC' +
                tiêu đề phụ lục in nghiêng)."""
    items = []
    pending_chuong_num = None
    pending_loai = None
    pending_phu_luc = False     # vừa gặp 'PHỤ LỤC' → paragraph kế là tiêu đề

    for kind_flow, obj in flow_items:
        # Bảng (Phụ lục) — phân biệt khối quốc hiệu của Mẫu biểu (không viền)
        # với bảng dữ liệu thường (có viền).
        if kind_flow == 't':
            pending_phu_luc = False
            if appendix and _is_header_table(obj):
                items.append(('header_block', obj))
            elif appendix and _is_sign_block(obj):
                items.append(('sign_block', obj))
            else:
                items.append(('table', obj))
            continue

        para = obj
        text = para.text.strip()
        if not text:
            continue

        # Đang chờ tên chương?
        if pending_chuong_num is not None:
            num, quote = pending_chuong_num
            items.append(('chuong', num, text, quote))
            pending_chuong_num = None
            continue

        # Đang chờ tiêu đề (sau loại văn bản)?
        if pending_loai is not None:
            items.append(('title', pending_loai, text))
            pending_loai = None
            continue

        # Đang chờ tiêu đề phụ lục (sau 'PHỤ LỤC')?
        if pending_phu_luc:
            title, ref = _split_phu_luc_title(para)
            if title:
                items.append(('phu_luc_title', title))
            if ref:
                items.append(('phu_luc_ref', ref))
            pending_phu_luc = False
            continue

        # Tiêu đề phụ lục
        if appendix and re.fullmatch(r'PHỤ\s+LỤC(\s+[IVXLCDM\d]+)?', text):
            items.append(('phu_luc', text))
            pending_phu_luc = True
            continue

        # Tiêu đề loại văn bản đính kèm (KẾ HOẠCH, CHƯƠNG TRÌNH, ĐỀ ÁN...)
        # ban hành kèm theo Quyết định — đứng riêng 1 dòng, căn giữa. Xử lý
        # như 'PHỤ LỤC': sang trang mới + tách tiêu đề phụ / dòng "(Kèm theo)".
        if text in ATTACHED_DOC_KEYWORDS and \
                para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            items.append(('attached_doc', text))
            pending_phu_luc = True
            continue

        kind = _classify(text, co_quan)
        tag = kind[0]

        if tag == 'loai_van_ban':
            pending_loai = kind[1]
        elif tag == 'chuong_num':
            pending_chuong_num = (kind[1], kind[2])
        elif tag in ('muc', 'dieu', 'can_cu', 'roman_section',
                     'co_quan_body', 'quyet_nghi', 'don_vi', 'mau_so'):
            items.append(kind)
        else:  # 'khoan'
            centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
            # Trong Phụ lục: giữ các dòng tiêu đề căn giữa của Mẫu biểu
            # (vd 'DANH SÁCH' đậm, '(Ban hành kèm theo...)' nghiêng).
            if appendix and centered and DANH_MUC_HEADING_RE.match(text):
                items.append(('danh_muc', text))
            elif appendix and centered and _is_para_bold(para):
                items.append(('center_bold', text))
            elif appendix and centered and _is_para_italic(para):
                items.append(('center_italic', text))
            # Trong Phụ lục: dòng căn giữa thường (không đậm/nghiêng) —
            # vd 'Kính gửi: ...', tiêu đề phụ 'CÁC MẪU BAN HÀNH KÈM THEO...'.
            # Giữ căn giữa thay vì hạ về đoạn thường (thụt dòng, căn đều).
            elif appendix and centered:
                items.append(('center_normal', text))
            # Trong thân bài: dòng đậm + căn giữa đứng riêng (chưa khớp loại
            # văn bản / câu lệnh ban hành / Chương / Mục ở trên) là dòng chức
            # danh người ban hành — vd "BỘ TRƯỞNG BỘ VĂN HÓA, THỂ THAO VÀ DU
            # LỊCH" trong Quyết định. Giữ đậm, căn giữa.
            elif not appendix and centered and _is_para_bold(para):
                items.append(('co_quan_body', text))
            # Kiểm tra inline italic / mixed / bold formatting
            elif _has_mixed_inline_format(para):
                items.append(('mixed', _runs_with_format(para)))
            elif _is_para_italic(para):
                items.append(('italic_body', text))
            elif appendix and _is_para_bold(para):
                # Tiêu đề mục đậm trong Phụ lục (vd '1. Các bộ...') — giữ đậm.
                # Trong thân bài chính KHÔNG giữ đậm các đoạn số thứ tự.
                items.append(('bold_body', text))
            else:
                items.append(kind)

    # Fallback các pending còn dở
    if pending_loai is not None:
        items.append(('title', pending_loai, ''))
    if pending_chuong_num is not None:
        num, quote = pending_chuong_num
        items.append(('chuong', num, '', quote))

    return items


def _find_appendix_start(flow, start):
    """Chỉ số phần tử (paragraph) bắt đầu phần Phụ lục — mốc 'PHỤ LỤC',
    'Mẫu số NN' đứng riêng, hoặc tiêu đề loại văn bản đính kèm căn giữa
    (KẾ HOẠCH, ĐỀ ÁN, QUY CHẾ...). Trả về None nếu không có phụ lục.

    Dùng để khoanh vùng tìm bảng kết (chữ ký) — bảng kết của VĂN BẢN nằm
    NGAY TRƯỚC phần phụ lục; các khối 'Nơi nhận' / chữ ký bên trong từng
    Mẫu biểu (nằm SAU mốc này) không được nhầm là bảng kết."""
    for i in range(start, len(flow)):
        kind, obj = flow[i]
        if kind != 'p':
            continue
        text = obj.text.strip()
        if not text:
            continue
        if re.fullmatch(r'PHỤ\s+LỤC(\s+[IVXLCDM\d]+)?', text):
            return i
        if re.fullmatch(r'Mẫu số\s+\S+', text):
            return i
        if text in ATTACHED_DOC_KEYWORDS and \
                obj.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return i
    return None


def _find_closing_table_index(flow, first_table_idx, upper=None):
    """Tìm chỉ số bảng kết (Nơi nhận / chữ ký) trong flow.

    Bảng kết là bảng CÓ chứa 'Nơi nhận' hoặc khối chức danh ký
    (TM., THỪA LỆNH, CHỦ TỊCH, BỘ TRƯỞNG, THỦ TƯỚNG...). Trả về None nếu
    không xác định được — khi đó coi như không có khối kết.

    upper: chỉ xét các bảng có chỉ số < upper (mốc bắt đầu Phụ lục). Nhờ
    đó tránh nhầm khối 'Nơi nhận' / chữ ký bên trong các Mẫu biểu của Phụ
    lục (vd Thông tư kèm nhiều Mẫu đơn) là bảng kết của văn bản.

    QUAN TRỌNG: phải phân biệt với bảng dữ liệu trong Phụ lục (vốn cũng
    có ≥ 2 cột). Vì vậy ưu tiên nhận dạng theo dấu hiệu chức danh ký;
    chỉ khi không tìm thấy mới fallback theo bảng 2 cột & ít hàng (≤ 3)
    — đặc trưng của khối ký, khác bảng dữ liệu nhiều hàng."""
    SIGN_HINTS = ('Nơi nhận', 'TM.', 'KT.', 'THỪA LỆNH', 'CHỦ TỊCH',
                  'BỘ TRƯỞNG', 'THỦ TƯỚNG', 'TUQ.', 'TỔNG THƯ KÝ',
                  'CHỦ NHIỆM', 'VIỆN TRƯỞNG', 'CHÁNH ÁN', 'CHỦ TỊCH NƯỚC')

    def cells_text(table):
        return ' '.join(c.text for row in table.rows for c in row.cells)

    hi = (upper if upper is not None else len(flow)) - 1

    # Pass 1: nhận dạng theo dấu hiệu chức danh ký (đáng tin cậy nhất)
    for i in range(hi, -1, -1):
        if flow[i][0] != 't' or i == first_table_idx:
            continue
        if any(h in cells_text(flow[i][1]) for h in SIGN_HINTS):
            return i

    # Pass 2 (fallback): bảng 2 cột & ≤ 3 hàng — dạng khối ký không có
    # dấu hiệu rõ; loại trừ bảng dữ liệu Phụ lục nhiều hàng.
    for i in range(hi, -1, -1):
        if flow[i][0] != 't' or i == first_table_idx:
            continue
        table = flow[i][1]
        if len(table.columns) == 2 and len(table.rows) <= 3:
            return i
    return None


def extract_body_items(doc, co_quan=''):
    """Đi qua body của document, trả về (items, appendix, closing_table).

    Giữ nguyên cả bảng dữ liệu trong thân bài & Phụ lục (interleave theo
    đúng thứ tự). Hỗ trợ phần Phụ lục nằm SAU bảng chữ ký (vd nghị quyết
    Quốc hội / Chính phủ kèm Phụ lục số liệu)."""
    body_xml = doc.element.body

    # Liệt kê tất cả paragraphs + tables theo đúng thứ tự xuất hiện
    flow = []   # list of ('p', Paragraph) | ('t', Table)
    for child in body_xml.iterchildren():
        if child.tag == qn('w:tbl'):
            from docx.table import Table  # noqa: import cục bộ
            flow.append(('t', Table(child, doc.part)))
        elif child.tag == qn('w:p'):
            flow.append(('p', Paragraph(child, doc.part)))

    table_idxs = [i for i, (k, _) in enumerate(flow) if k == 't']
    if not table_idxs:
        raise ValueError("Không tìm thấy bảng quốc hiệu.")

    # Quốc hiệu = bảng đầu tiên
    first_table_idx = table_idxs[0]
    start = first_table_idx + 1

    # Mốc bắt đầu Phụ lục — khoanh vùng tìm bảng kết (tránh nhầm các khối
    # 'Nơi nhận' / chữ ký bên trong Mẫu biểu là bảng kết của văn bản).
    appendix_start = _find_appendix_start(flow, start)

    # Bảng kết (chữ ký / Nơi nhận) ngăn thân bài & phụ lục — chỉ tìm trong
    # phần TRƯỚC mốc phụ lục.
    closing_idx = _find_closing_table_index(
        flow, first_table_idx, upper=appendix_start)
    closing_table = None
    if closing_idx is not None and closing_idx > first_table_idx:
        # Có bảng kết: thân bài trước nó, phụ lục là toàn bộ phần sau.
        closing_table = flow[closing_idx][1]
        main_flow = flow[start:closing_idx]
        appendix_flow = flow[closing_idx + 1:]
    elif appendix_start is not None:
        # Không có bảng kết nhưng có phụ lục (vd văn bản ký ngay trong mẫu):
        # cắt tại mốc phụ lục, không có khối kết riêng.
        main_flow = flow[start:appendix_start]
        appendix_flow = flow[appendix_start:]
    else:
        # Không có bảng kết lẫn phụ lục → toàn bộ phần sau quốc hiệu là thân bài
        main_flow = flow[start:]
        appendix_flow = []

    items = _classify_flow(main_flow, co_quan=co_quan)
    appendix = _classify_flow_appendix(appendix_flow, co_quan=co_quan)

    return items, appendix, closing_table


def _classify_flow_appendix(appendix_flow, co_quan=''):
    """Phân loại phần Phụ lục, tách riêng các Mẫu biểu.

    Phần TRƯỚC Mẫu đầu tiên (tiêu đề PHỤ LỤC + nội dung trình tự, thủ tục…)
    được dựng lại theo style chuẩn như cũ. Từ Mẫu đầu tiên trở đi, toàn bộ
    các Mẫu + nội dung được GIỮ NGUYÊN định dạng gốc (chỉ chuẩn hóa font +
    cỡ chữ) — trả về dưới dạng item thô ('mau_p' / 'mau_t')."""
    mau_start = None
    for i, (kind, obj) in enumerate(appendix_flow):
        if kind == 'p' and MAU_HEADING_RE.match(obj.text.strip()):
            mau_start = i
            break

    if mau_start is None:
        return _classify_flow(appendix_flow, co_quan=co_quan, appendix=True)

    pre = _classify_flow(appendix_flow[:mau_start],
                         co_quan=co_quan, appendix=True)
    mau_items = [('mau_t', obj) if kind == 't' else ('mau_p', obj)
                 for kind, obj in appendix_flow[mau_start:]]
    return pre + mau_items


def extract_closing(table):
    """Trích Nơi nhận, chức danh, người ký từ bảng kết đã xác định.

    table: bảng kết (chữ ký / Nơi nhận) do extract_body_items tìm ra,
    hoặc None nếu văn bản không có khối kết."""
    if table is None:
        return None
    if not table.rows or len(table.rows[0].cells) < 2:
        return None

    # Ô trái: danh sách nơi nhận
    left = table.rows[0].cells[0]
    left_text = '\n'.join(p.text for p in left.paragraphs)

    # Chỉ tách theo newline — KHÔNG tách trên ' - ' inline
    # (vì dấu ' - ' có thể xuất hiện trong tên nơi nhận, vd:
    # 'tổ chức chính trị - xã hội')
    recipients = []
    for line in left_text.split('\n'):
        line = line.strip()
        if not line or 'Nơi nhận' in line:
            continue
        # Bỏ dấu '-' đầu (cũng chấp nhận trường hợp thiếu khoảng trắng: "-Văn phòng...")
        line = re.sub(r'^[-–—]\s*', '', line).strip()
        # Bỏ dấu ';' hoặc '.' cuối
        line = re.sub(r'[;.]\s*$', '', line).strip()
        # Bỏ ký tự  no-break space
        line = line.replace('\xa0', ' ')
        # Gộp nhiều khoảng trắng
        line = re.sub(r'\s{2,}', ' ', line)
        if line:
            recipients.append(line)

    # Ô phải: chức danh (in hoa) + chữ ký + tên người ký
    right = table.rows[0].cells[1]
    # Tách theo cả paragraph và line break
    right_lines = []
    for p in right.paragraphs:
        for line in p.text.split('\n'):
            line = line.strip()
            if line:
                right_lines.append(line)

    position_lines = []
    signer = ''
    for line in right_lines:
        # In hoa toàn bộ ⇒ chức danh
        if line.upper() == line and len(line) > 1:
            position_lines.append(line)
        else:
            # Dòng có chữ thường ⇒ tên người ký (lấy dòng đầu tiên)
            if not signer:
                signer = line

    return {
        'recipients': recipients,
        'position_lines': position_lines,
        'signer': signer,
    }


# ============================================================
# DỰNG LẠI VĂN BẢN — Tạo document mới với định dạng chuẩn
# ============================================================

def _add_header_table(doc, co_quan, so_van_ban, dia_diem,
                      ngay_ban_hanh, prefix_so):
    table = doc.add_table(rows=2, cols=2)
    _set_table_borders_none(table)
    # Chia 2 cột theo tỷ lệ 36/64 (cơ quan / quốc hiệu) trên bề rộng dùng
    # được ~16.5cm. Đặt cả width của cột lẫn từng ô + tắt autofit để Word
    # giữ đúng tỷ lệ (chỉ đặt columns[].width thường không được tôn trọng).
    COL_LEFT = Cm(5.94)    # 36%
    COL_RIGHT = Cm(10.56)  # 64%
    table.allow_autofit = False
    table.columns[0].width = COL_LEFT
    table.columns[1].width = COL_RIGHT

    # Hàng 1 trái: tên cơ quan
    left = table.cell(0, 0)
    left.width = COL_LEFT
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(left.paragraphs[0], co_quan, bold=True, size=SIZE_DEFAULT)
    p = left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, '--------', bold=True, size=SIZE_DEFAULT)

    # Hàng 1 phải: quốc hiệu
    right = table.cell(0, 1)
    right.width = COL_RIGHT
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(right.paragraphs[0],
             'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM',
             bold=True, size=SIZE_DEFAULT)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, 'Độc lập - Tự do - Hạnh phúc', bold=True, size=SIZE_DEFAULT)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, '---------------', bold=True, size=SIZE_DEFAULT)

    # Hàng 2 trái: số văn bản
    left2 = table.cell(1, 0)
    left2.width = COL_LEFT
    left2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(left2.paragraphs[0], prefix_so + so_van_ban, size=SIZE_DOC_NUMBER)

    # Hàng 2 phải: địa điểm, ngày
    right2 = table.cell(1, 1)
    right2.width = COL_RIGHT
    right2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(right2.paragraphs[0],
             f'{dia_diem}, ngày {ngay_ban_hanh}',
             italic=True, size=SIZE_BODY)


def _compact(p):
    """Giãn dòng đơn (single) + bỏ cách đoạn sau cho 1 đoạn thân bài —
    chuẩn trình bày VBHC Việt Nam. Tránh giá trị mặc định của python-docx
    (giãn dòng 1.15, cách đoạn sau ~10pt) làm văn bản bị thưa, dài thêm trang.
    Chỉ áp cho các đoạn DỰNG LẠI; không đụng tới khối tiêu đề, bảng, hay Mẫu
    biểu giữ nguyên gốc (để các phần đó kế thừa mặc định như mẫu desired)."""
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = SPACE_ZERO
    return p


def _add_title_block(doc, loai_van_ban, tieu_de):
    # Đoạn trống
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, ' ', bold=True, size=SIZE_DEFAULT)

    # Loại văn bản
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, loai_van_ban, bold=True, size=SIZE_DEFAULT)

    # Tiêu đề dài
    if tieu_de:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = SPACE_BEFORE
        _add_run(p, tieu_de, bold=True, size=SIZE_HEADING)


def _add_can_cu(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, italic=True, size=SIZE_BODY)


def _add_co_quan_body(doc, text):
    """Dòng cơ quan ban hành đứng riêng (vd QUỐC HỘI, CHÍNH PHỦ) —
    đậm, căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text.upper(), bold=True, size=SIZE_DEFAULT)


def _add_quyet_nghi(doc, text):
    """Câu lệnh ban hành (QUYẾT NGHỊ:, NGHỊ ĐỊNH:...) — đậm, căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, bold=True, size=SIZE_BODY)


def _add_roman_section(doc, roman, name):
    """Tiêu đề mục La Mã kiểu Chính phủ (I., II....): đậm, căn đều,
    thụt dòng đầu, tiêu đề cùng dòng."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, f'{roman}. {name}', bold=True, size=SIZE_HEADING)


def _add_bold_body(doc, text):
    """Đoạn thân bài đậm (vd tiêu đề '1. Quán triệt...' trong NQ-CP):
    đậm, căn đều, thụt dòng đầu."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, bold=True, size=SIZE_BODY)


def _add_phu_luc(doc, text):
    """Tiêu đề 'PHỤ LỤC' — đậm, căn giữa, sang trang mới."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text.upper(), bold=True, size=SIZE_HEADING)


def _add_attached_doc(doc, text):
    """Tiêu đề loại văn bản đính kèm (KẾ HOẠCH, ĐỀ ÁN...) — đậm, căn giữa,
    KHÔNG sang trang mới (khác 'PHỤ LỤC')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text.upper(), bold=True, size=SIZE_BODY)


def _add_phu_luc_title(doc, text):
    """Tiêu đề phụ của phụ lục — đậm, căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, bold=True, size=SIZE_HEADING)


def _add_phu_luc_ref(doc, text):
    """Dòng '(Kèm theo Nghị quyết số ... )' dưới tiêu đề phụ lục —
    in nghiêng, căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, italic=True, size=SIZE_BODY)


def _add_mau_so(doc, text):
    """Tiêu đề 'Mẫu số 01' của một mẫu biểu trong Phụ lục — đậm, căn phải.
    Các mẫu chạy liền nhau (không ngắt trang) theo mẫu desired ND_so_132."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, bold=True, size=SIZE_BODY)


def _add_danh_muc(doc, text):
    """Tiêu đề 'DANH MỤC N' của một danh mục mẫu biểu trong Phụ lục — đậm,
    căn giữa, sang trang mới (mỗi danh mục bắt đầu một trang riêng)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text.upper(), bold=True, size=SIZE_BODY)


def _add_center_bold(doc, text):
    """Dòng tiêu đề căn giữa, đậm trong mẫu biểu (vd 'DANH SÁCH',
    'Công nhận người có uy tín...')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, bold=True, size=SIZE_BODY)


def _add_center_italic(doc, text):
    """Dòng căn giữa, in nghiêng trong mẫu biểu (vd '(Ban hành kèm theo
    Quyết định số ...)')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, italic=True, size=SIZE_BODY)


def _add_header_block(doc, src_table):
    """Dựng lại khối quốc hiệu của một mẫu biểu (cơ quan trái / CỘNG HÒA...
    phải) — KHÔNG viền, căn giữa, giữ định dạng đậm/nghiêng của từng ô."""
    n_rows = len(src_table.rows)
    n_cols = len(src_table.columns)
    if n_rows == 0 or n_cols == 0:
        return
    table = doc.add_table(rows=n_rows, cols=n_cols)
    _set_table_borders_none(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = _source_col_widths(src_table)
    use_src_w = len(col_w) == n_cols and all(w for w in col_w)
    if use_src_w:
        table.allow_autofit = False
        for i, w in enumerate(col_w):
            table.columns[i].width = Pt(w / 20)

    for r, row in enumerate(src_table.rows):
        _set_cant_split(table.rows[r])
        src_cells = row.cells
        for c in range(n_cols):
            src_cell = src_cells[c] if c < len(src_cells) else None
            dest = table.cell(r, c)
            if use_src_w:
                dest.width = Pt(col_w[c] / 20)
            p = dest.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = SPACE_ZERO
            p.paragraph_format.space_after = SPACE_ZERO
            if src_cell is not None:
                _copy_cell_runs(p, src_cell, size=SIZE_DEFAULT)


def _add_sign_block(doc, src_table):
    """Dựng lại khối chữ ký / Nơi nhận bên trong một Mẫu biểu (vd
    'NGƯỜI ĐẠI DIỆN THEO PHÁP LUẬT', 'THỦ TRƯỞNG CƠ QUAN', 'Người lập
    báo cáo') — KHÔNG viền (khác bảng dữ liệu), giữ định dạng đậm/nghiêng
    của từng run. Ô chứa 'Nơi nhận' căn trái + cỡ chữ nhỏ; các ô còn lại
    (khối chức danh + chữ ký) căn giữa."""
    n_rows = len(src_table.rows)
    n_cols = len(src_table.columns)
    if n_rows == 0 or n_cols == 0:
        return
    table = doc.add_table(rows=n_rows, cols=n_cols)
    _set_table_borders_none(table)

    col_w = _source_col_widths(src_table)
    use_src_w = len(col_w) == n_cols and all(w for w in col_w)
    if use_src_w:
        table.allow_autofit = False
        for i, w in enumerate(col_w):
            table.columns[i].width = Pt(w / 20)

    for r, row in enumerate(src_table.rows):
        _set_cant_split(table.rows[r])
        src_cells = row.cells
        for c in range(n_cols):
            src_cell = src_cells[c] if c < len(src_cells) else None
            dest = table.cell(r, c)
            if use_src_w:
                dest.width = Pt(col_w[c] / 20)
            p = dest.paragraphs[0]
            p.paragraph_format.space_before = SPACE_ZERO
            p.paragraph_format.space_after = SPACE_ZERO
            if src_cell is None:
                continue
            noi_nhan = 'Nơi nhận' in src_cell.text
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if noi_nhan
                           else WD_ALIGN_PARAGRAPH.CENTER)
            _copy_cell_runs(
                p, src_cell,
                size=SIZE_RECIPIENTS if noi_nhan else SIZE_BODY)


def _add_center_normal(doc, text):
    """Dòng căn giữa, không đậm/nghiêng trong mẫu biểu (vd 'Kính gửi: ...',
    tiêu đề phụ 'CÁC MẪU BAN HÀNH KÈM THEO PHỤ LỤC')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, size=SIZE_BODY)


def _add_don_vi(doc, text):
    """Dòng đơn vị tính của bảng ('Đơn vị: Tỷ đồng') — in nghiêng,
    căn phải, không thụt dòng đầu."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, italic=True, size=SIZE_BODY)


def _cell_text_lines(cell):
    """Lấy text của cell, giữ ngắt dòng giữa các paragraph."""
    return '\n'.join(p.text for p in cell.paragraphs).strip()


def _looks_numeric(text):
    """Cell chỉ gồm số (kèm . , % - khoảng trắng) → căn phải."""
    t = text.strip()
    if not t:
        return False
    return bool(re.fullmatch(r'[\d.,%()\-\s]+', t)) and any(c.isdigit() for c in t)


def _set_cell_margins(cell, left, right):
    """Đặt lề trái/phải của cell (đơn vị: twip — 1cm ≈ 567 twip)."""
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for name, val in (('left', left), ('right', right)):
        m = OxmlElement(f'w:{name}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _source_col_widths(src_table):
    """Bề rộng các cột (twip) lấy từ <w:tblGrid> của bảng nguồn.
    Trả về [] nếu không có grid hợp lệ."""
    grid = src_table._tbl.find(qn('w:tblGrid'))
    if grid is None:
        return []
    widths = []
    for col in grid.findall(qn('w:gridCol')):
        w = col.get(qn('w:w'))
        widths.append(int(w) if (w and w.isdigit()) else None)
    return widths


def _grid_span(tc):
    """Số cột mà ô <w:tc> chiếm ngang (gridSpan) — mặc định 1."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return 1
    gs = tcPr.find(qn('w:gridSpan'))
    if gs is None:
        return 1
    v = gs.get(qn('w:val'))
    return int(v) if (v and v.isdigit()) else 1


def _copy_row_height(dest_row, src_row):
    """Giữ nguyên chiều cao hàng từ bảng nguồn (nếu có trHeight).
    Hàng không khai báo chiều cao ⇒ để tự co theo nội dung (giữ nguyên)."""
    trPr = src_row._tr.find(qn('w:trPr'))
    h = trPr.find(qn('w:trHeight')) if trPr is not None else None
    if h is None:
        return
    val = h.get(qn('w:val'))
    if not (val and val.isdigit()):
        return
    dest_row.height = Pt(int(val) / 20)   # twip → point
    rule = h.get(qn('w:hRule'))
    dest_row.height_rule = {
        'exact': WD_ROW_HEIGHT_RULE.EXACTLY,
        'atLeast': WD_ROW_HEIGHT_RULE.AT_LEAST,
        'auto': WD_ROW_HEIGHT_RULE.AUTO,
    }.get(rule, WD_ROW_HEIGHT_RULE.AT_LEAST)


def _set_table_cell_margins(table, *, top, bottom=0, left=108, right=108):
    """Đặt lề mặc định cho MỌI ô của bảng (đơn vị twip). Dùng để thêm
    đệm phía trên mỗi hàng cho bảng đỡ chật. Ô có tcMar riêng chỉ ghi
    đè theo từng cạnh nên top vẫn được kế thừa từ đây."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblCellMar'))
    if old is not None:
        tblPr.remove(old)
    mar = OxmlElement('w:tblCellMar')
    for name, val in (('top', top), ('left', left),
                      ('bottom', bottom), ('right', right)):
        m = OxmlElement(f'w:{name}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    # tblCellMar đứng trước tblLook trong schema
    _AFTER = {qn('w:tblLook'), qn('w:tblCaption'), qn('w:tblDescription')}
    insert_at = len(tblPr)
    for i, child in enumerate(tblPr):
        if child.tag in _AFTER:
            insert_at = i
            break
    tblPr.insert(insert_at, mar)


def _copy_cell_runs(dest_para, src_cell, *, force_bold=None, size=SIZE_TABLE):
    """Chép nội dung 1 cell nguồn sang paragraph đích, GIỮ định dạng
    in đậm / in nghiêng của từng run. Ngắt dòng giữa các paragraph nguồn
    được giữ bằng <w:br/>.

    force_bold: True/False để ép đậm (vd hàng tiêu đề); None = giữ nguyên.
    """
    first = True
    for sp in src_cell.paragraphs:
        runs = [r for r in sp.runs if r.text]
        if not first and runs:
            dest_para.add_run().add_break()
        for r in runs:
            bold = force_bold if force_bold is not None else bool(r.bold)
            _add_run(dest_para, r.text, bold=bold,
                     italic=bool(r.italic), size=size)
        if runs:
            first = False


def _add_data_table(doc, src_table):
    """Dựng lại bảng dữ liệu (Phụ lục) có viền đơn, Times New Roman,
    giữ định dạng đậm/nghiêng của từng ô.

    - Hàng tiêu đề: đậm, căn giữa ngang + dọc.
    - Cột đầu là STT/TT: căn giữa ngang + dọc, lề ngang nhỏ hai bên.
    - Ô số: căn phải. Ô còn lại cột đầu: căn giữa. Còn lại: căn trái.
    - Giữ nguyên bề rộng cột & chiều cao hàng của bảng nguồn.
    """
    n_rows = len(src_table.rows)
    n_cols = len(src_table.columns)
    if n_rows == 0 or n_cols == 0:
        return

    # Cột đầu có phải STT/TT không?
    header0 = _cell_text_lines(src_table.rows[0].cells[0]).strip().upper()
    stt_col = header0 in ('STT', 'TT', 'SỐ TT', 'SỐ THỨ TỰ')

    # Tiêu đề có thể trải trên NHIỀU hàng (vd bảng chấm điểm: hàng 1 gộp
    # "Điểm" / "Kết quả đánh giá", hàng 2 là tiêu đề con "Tối đa | Tối
    # thiểu" / "Đạt | Không đạt"). Khi cột STT/TT gộp dọc, các hàng tiêu đề
    # con có ô đầu RỖNG → coi là hàng tiêu đề tiếp theo (đậm, căn giữa).
    # Dừng ở hàng đầu tiên có ô đầu khác rỗng (hàng dữ liệu hoặc hàng nhóm).
    header_rows = {0}
    if stt_col:
        from docx.table import _Cell   # noqa: import cục bộ
        for r in range(1, n_rows):
            # Dùng ô <w:tc> THÔ đầu hàng — không dùng rows[r].cells[0] vì
            # python-docx nhân bản nội dung ô gộp dọc (trả về 'TT' thay vì rỗng).
            tcs = src_table.rows[r]._tr.findall(qn('w:tc'))
            first_text = _cell_text_lines(_Cell(tcs[0], src_table)).strip() if tcs else ''
            if first_text:
                break
            header_rows.add(r)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    _set_table_borders_single(table)
    _set_table_cell_margins(table, top=TABLE_ROW_TOP_PAD)  # đệm trên mỗi hàng
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Giữ nguyên bề rộng cột của bảng nguồn (twip → point, layout cố định)
    col_w = _source_col_widths(src_table)
    use_src_w = len(col_w) == n_cols and all(w for w in col_w)
    if use_src_w:
        table.allow_autofit = False
        for i, w in enumerate(col_w):
            table.columns[i].width = Pt(w / 20)

    from docx.table import _Cell   # noqa: import cục bộ

    for r, row in enumerate(src_table.rows):
        _set_cant_split(table.rows[r])
        _copy_row_height(table.rows[r], row)   # giữ chiều cao hàng gốc
        is_header = r in header_rows

        # Đi theo từng ô <w:tc> nguồn (không dùng row.cells vì python-docx
        # nhân bản ô đã gộp ngang) để bảo toàn gridSpan — các dòng tiêu đề
        # nhóm (vd "I | THỦ TỤC HÀNH CHÍNH CẤP TỈNH") gộp 2 cột cuối.
        g = 0
        for tc in row._tr.findall(qn('w:tc')):
            span = _grid_span(tc)
            start = g
            end = min(g + span - 1, n_cols - 1)
            g += span
            if start >= n_cols:
                break

            src_cell = _Cell(tc, src_table)
            text = _cell_text_lines(src_cell)

            dest = table.cell(r, start)
            if end > start:
                dest = dest.merge(table.cell(r, end))
            if use_src_w:
                dest.width = Pt(sum(col_w[start:end + 1]) / 20)

            p = dest.paragraphs[0]
            p.paragraph_format.space_before = SPACE_ZERO
            p.paragraph_format.space_after = SPACE_ZERO

            if is_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dest.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            elif start == 0 and stt_col:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dest.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            elif end > start:
                # Ô gộp ngang = dòng tiêu đề nhóm → căn trái, giữ đậm gốc
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif _looks_numeric(text):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif start == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Cột STT: lề ngang nhỏ hai bên
            if start == 0 and stt_col:
                _set_cell_margins(dest, 28, 28)   # ~0.05cm mỗi bên

            force_bold = True if is_header else None
            _copy_cell_runs(p, src_cell, force_bold=force_bold,
                            size=SIZE_TABLE)


def _add_chuong(doc, roman, name, quote=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, f'{quote}Chương {roman}', bold=True, size=SIZE_HEADING)

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = SPACE_BEFORE
        _compact(p)
        _add_run(p, name.upper(), bold=True, size=SIZE_HEADING)


def _add_muc(doc, so, name, quote=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, f'{quote}Mục {so}. {name.upper()}', bold=True, size=SIZE_HEADING)


def _add_dieu(doc, so, name, quote=''):
    """Tiêu đề Điều — đậm, căn đều, thụt dòng đầu.

    quote: dấu ngoặc kép mở (vd “) nếu Điều này là phần trích dẫn được
    chèn/sửa đổi trong văn bản sửa đổi, bổ sung (Điều nằm trong Điều) —
    giữ lại để bảo toàn ý nghĩa pháp lý."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, f'{quote}Điều {so}. {name}', bold=True, size=SIZE_BODY)


def _add_khoan(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, size=SIZE_BODY)


def _add_italic_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    _add_run(p, text, italic=True, size=SIZE_BODY)


def _add_mixed(doc, parts):
    """Đoạn body với nhiều run định dạng khác nhau (giữ inline italic/bold)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT_FIRST
    p.paragraph_format.space_before = SPACE_BEFORE
    _compact(p)
    for text, fmt in parts:
        _add_run(p, text,
                 bold=fmt.get('bold', False),
                 italic=fmt.get('italic', False),
                 size=SIZE_BODY)


def _add_closing_table(doc, recipients, position_lines, signer,
                       blank_lines=5):
    # Đoạn trống ngăn cách
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, ' ', size=SIZE_DEFAULT)

    table = doc.add_table(rows=1, cols=2)
    _set_table_borders_none(table)
    _set_cant_split(table.rows[0])
    table.columns[0].width = Cm(8.25)
    table.columns[1].width = Cm(8.25)

    # Ô trái: Nơi nhận
    left = table.cell(0, 0)
    p = left.paragraphs[0]
    p.paragraph_format.space_before = SPACE_ZERO
    p.paragraph_format.space_after = SPACE_ZERO
    if recipients:
        _add_run(p, 'Nơi nhận:', bold=True, italic=True, size=SIZE_RECIPIENTS)
        for line in recipients:
            p = left.add_paragraph()
            p.paragraph_format.space_before = SPACE_ZERO
            p.paragraph_format.space_after = SPACE_ZERO
            text = line.strip()
            if not text.startswith('-'):
                text = '- ' + text
            if not text.endswith(';') and not text.endswith('.'):
                text = text + ';'
            _add_run(p, text, size=SIZE_RECIPIENTS)

    # Ô phải: chức danh + chữ ký
    right = table.cell(0, 1)
    first = True
    for line in position_lines:
        p = right.paragraphs[0] if first else right.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = SPACE_ZERO
        p.paragraph_format.space_after = SPACE_ZERO
        _add_run(p, line, bold=True, size=SIZE_BODY)
    # Dòng trống chừa cho chữ ký
    for _ in range(blank_lines):
        p = right.add_paragraph() if not first else right.paragraphs[0]
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = SPACE_ZERO
        p.paragraph_format.space_after = SPACE_ZERO
        _add_run(p, ' ', size=SIZE_BODY)
    # Tên người ký
    if signer:
        p = right.add_paragraph() if not first else right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = SPACE_ZERO
        p.paragraph_format.space_after = SPACE_ZERO
        _add_run(p, signer, bold=True, size=SIZE_BODY)


def _force_run_font(run, size):
    """Ép font + cỡ chữ cho 1 run, GIỮ nguyên các định dạng khác
    (đậm/nghiêng/gạch chân/căn lề…). Dùng khi chép Mẫu biểu nguyên trạng."""
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), DEFAULT_FONT)


def _append_raw_element(doc, src_elem):
    """Chép nguyên (deep copy) 1 phần tử XML (paragraph/table) của document
    gốc vào CUỐI thân document mới — chèn TRƯỚC <w:sectPr> để giữ phần thiết
    lập trang. Giữ nguyên toàn bộ định dạng gốc của phần tử."""
    new_elem = copy.deepcopy(src_elem)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(new_elem)
    else:
        body.append(new_elem)
    return new_elem


def _add_mau_paragraph(doc, src_para):
    """Chép nguyên 1 paragraph của Mẫu biểu — giữ định dạng gốc (căn lề,
    thụt dòng, đậm/nghiêng, dòng chấm…), chỉ chuẩn hóa font + cỡ chữ."""
    new_p = _append_raw_element(doc, src_para._p)
    para = Paragraph(new_p, doc.part)
    # 'DANH MỤC N' (danh mục mẫu biểu) vẫn bắt đầu một trang mới. Còn các
    # 'Mẫu số NN' chạy liền nhau (KHÔNG ngắt trang) — theo mẫu desired
    # ND_so_132, các mẫu chỉ cách nhau bằng dòng trống của file gốc.
    t = para.text.strip()
    if DANH_MUC_HEADING_RE.match(t):
        para.paragraph_format.page_break_before = True
    for run in para.runs:
        _force_run_font(run, SIZE_MAU)


def _add_mau_table(doc, src_table):
    """Chép nguyên 1 bảng của Mẫu biểu — giữ viền, bề rộng cột, ô gộp, căn
    lề… của bảng gốc, chỉ chuẩn hóa font + cỡ chữ trong mọi ô."""
    from docx.table import Table  # noqa: import cục bộ
    new_tbl = _append_raw_element(doc, src_table._tbl)
    for row in Table(new_tbl, doc.part).rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    _force_run_font(run, SIZE_MAU)


def _make_dispatch(doc):
    """Bảng điều phối tag → hàm dựng đoạn, dùng chung cho body & phụ lục."""
    return {
        'title':         lambda it: _add_title_block(doc, it[1], it[2]),
        'can_cu':        lambda it: _add_can_cu(doc, it[1]),
        'co_quan_body':  lambda it: _add_co_quan_body(doc, it[1]),
        'quyet_nghi':    lambda it: _add_quyet_nghi(doc, it[1]),
        'chuong':        lambda it: _add_chuong(doc, it[1], it[2],
                                                it[3] if len(it) > 3 else ''),
        'muc':           lambda it: _add_muc(doc, it[1], it[2],
                                             it[3] if len(it) > 3 else ''),
        'dieu':          lambda it: _add_dieu(doc, it[1], it[2],
                                              it[3] if len(it) > 3 else ''),
        'roman_section': lambda it: _add_roman_section(doc, it[1], it[2]),
        'khoan':         lambda it: _add_khoan(doc, it[1]),
        'italic_body':   lambda it: _add_italic_body(doc, it[1]),
        'bold_body':     lambda it: _add_bold_body(doc, it[1]),
        'mixed':         lambda it: _add_mixed(doc, it[1]),
        'phu_luc':       lambda it: _add_phu_luc(doc, it[1]),
        'attached_doc':  lambda it: _add_attached_doc(doc, it[1]),
        'phu_luc_title': lambda it: _add_phu_luc_title(doc, it[1]),
        'phu_luc_ref':   lambda it: _add_phu_luc_ref(doc, it[1]),
        'mau_so':        lambda it: _add_mau_so(doc, it[1]),
        'danh_muc':      lambda it: _add_danh_muc(doc, it[1]),
        'center_bold':   lambda it: _add_center_bold(doc, it[1]),
        'center_italic': lambda it: _add_center_italic(doc, it[1]),
        'center_normal': lambda it: _add_center_normal(doc, it[1]),
        'don_vi':        lambda it: _add_don_vi(doc, it[1]),
        'table':         lambda it: _add_data_table(doc, it[1]),
        'header_block':  lambda it: _add_header_block(doc, it[1]),
        'sign_block':    lambda it: _add_sign_block(doc, it[1]),
        'mau_p':         lambda it: _add_mau_paragraph(doc, it[1]),
        'mau_t':         lambda it: _add_mau_table(doc, it[1]),
    }


def build_document(metadata, body_items, closing, appendix=None):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_OTHER
    section.right_margin = MARGIN_OTHER
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    _set_title_page(section)

    # Page number ở header
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(p)

    _set_default_font(doc, DEFAULT_FONT, SIZE_DEFAULT)

    # Quốc hiệu
    _add_header_table(doc, **metadata)

    # Body
    dispatch = _make_dispatch(doc)
    for item in body_items:
        handler = dispatch.get(item[0])
        if handler:
            handler(item)

    # Closing
    if closing and (closing['recipients'] or closing['position_lines']
                    or closing['signer']):
        _add_closing_table(doc, **closing)

    # Phụ lục (sau khối chữ ký)
    if appendix:
        for item in appendix:
            handler = dispatch.get(item[0])
            if handler:
                handler(item)

    # Vá lỗi zoom của python-docx
    zoom = doc.settings.element.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')

    return doc


# ============================================================
# MAIN
# ============================================================

def reformat(input_path, output_path, verbose=True):
    """Đọc input docx → trích xuất → dựng lại theo chuẩn → lưu output."""
    if verbose:
        print(f"📖 Đọc: {input_path}")
    src = Document(str(input_path))

    metadata = extract_metadata(src)
    body_items, appendix, closing_table = extract_body_items(
        src, co_quan=metadata['co_quan'])
    closing = extract_closing(closing_table)

    if verbose:
        print(f"   ▸ Cơ quan       : {metadata['co_quan']}")
        print(f"   ▸ Số văn bản    : {metadata['prefix_so']}{metadata['so_van_ban']}")
        print(f"   ▸ Ngày ban hành : {metadata['ngay_ban_hanh']}")
        # Đếm các loại item
        counts = {}
        for it in body_items + appendix:
            counts[it[0]] = counts.get(it[0], 0) + 1
        print(f"   ▸ Nội dung phát hiện:")
        labels = {
            'title': 'Tiêu đề', 'can_cu': 'Căn cứ', 'chuong': 'Chương',
            'muc': 'Mục', 'dieu': 'Điều', 'khoan': 'Khoản/đoạn',
            'italic_body': 'Đoạn in nghiêng', 'mixed': 'Đoạn pha định dạng',
            'co_quan_body': 'Cơ quan ban hành', 'quyet_nghi': 'Câu lệnh ban hành',
            'roman_section': 'Mục La Mã (I, II...)', 'bold_body': 'Đoạn đậm',
            'phu_luc': 'Phụ lục', 'phu_luc_title': 'Tiêu đề phụ lục',
            'phu_luc_ref': 'Dòng "Kèm theo"', 'don_vi': 'Đơn vị tính',
            'table': 'Bảng dữ liệu', 'mau_so': 'Mẫu số',
            'center_bold': 'Tiêu đề mẫu (đậm)',
            'center_italic': 'Dòng mẫu (nghiêng)',
            'center_normal': 'Dòng mẫu (giữa)',
            'header_block': 'Khối quốc hiệu mẫu',
            'sign_block': 'Khối chữ ký mẫu',
            'mau_p': 'Đoạn Mẫu (giữ gốc)',
            'mau_t': 'Bảng Mẫu (giữ gốc)',
        }
        for tag, count in counts.items():
            print(f"       - {labels.get(tag, tag):20s}: {count}")
        if closing:
            print(f"   ▸ Khối kết      : {len(closing['recipients'])} nơi nhận, "
                  f"người ký: {closing['signer'] or '(không)'}")
        else:
            print(f"   ▸ Khối kết      : không có")
        if appendix:
            print(f"   ▸ Phụ lục       : {len(appendix)} đoạn")

    doc = build_document(metadata, body_items, closing, appendix=appendix)
    doc.save(str(output_path))
    if verbose:
        print(f"✅ Đã xuất: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='Định dạng lại văn bản pháp luật Việt Nam (.docx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Ví dụ:\n"
            "  python vbpl_reformat.py 'Nghị định 141.docx'\n"
            "  python vbpl_reformat.py input.docx output.docx\n"
        ),
    )
    parser.add_argument('input', help='Đường dẫn file .docx cần định dạng')
    parser.add_argument('output', nargs='?',
                        help='File xuất (mặc định: <input>_formatted.docx)')
    parser.add_argument('-q', '--quiet', action='store_true',
                        help='Không in thông tin chi tiết')
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"❌ Không tìm thấy file: {input_path}", file=sys.stderr)
        sys.exit(1)
    if input_path.suffix.lower() != '.docx':
        print(f"⚠️  Cảnh báo: File không có đuôi .docx",
              file=sys.stderr)

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = (input_path.parent /
                       f"{input_path.stem}_formatted.docx")

    try:
        reformat(input_path, output_path, verbose=not args.quiet)
    except PackageNotFoundError:
        print("❌ File bị hỏng hoặc không phải định dạng Word (.docx) hợp lệ.",
              file=sys.stderr)
        sys.exit(2)
    except ValueError as e:
        print(f"❌ Lỗi khi xử lý: {e}", file=sys.stderr)
        sys.exit(2)
    except Exception as e:
        print(f"❌ Lỗi không xác định: {e}", file=sys.stderr)
        raise


if __name__ == '__main__':
    main()
