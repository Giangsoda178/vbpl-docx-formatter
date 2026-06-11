"""Deep checks for the Vercel handler beyond test_api.mjs: the returned file
must be a python-docx-parseable document, and a structurally-unrecognized
docx (no national-emblem header table) must yield 422.

Usage: with serve_api_local.py running, python test/verify_vercel_handler.py [port]
"""

import io
import json
import sys
import urllib.request
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
BASE = f'http://127.0.0.1:{sys.argv[1] if len(sys.argv) > 1 else 3002}'
BOUNDARY = 'testboundary7423'
failures = 0


def post_docx(data, filename):
    fn = filename.encode('utf-8').decode('latin-1')  # raw UTF-8 bytes, as browsers send
    body = (
        f'--{BOUNDARY}\r\n'
        f'Content-Disposition: form-data; name="file"; filename="{fn}"\r\n'
        'Content-Type: application/octet-stream\r\n\r\n'
    ).encode('latin-1') + data + f'\r\n--{BOUNDARY}--\r\n'.encode('latin-1')
    req = urllib.request.Request(
        f'{BASE}/api/format', data=body, method='POST',
        headers={'Content-Type': f'multipart/form-data; boundary={BOUNDARY}'},
    )
    try:
        with urllib.request.urlopen(req) as res:
            return res.status, res.read()
    except urllib.error.HTTPError as e:
        return e.code, e.read()


def check(label, cond, detail=''):
    global failures
    print(f'{"PASS" if cond else "FAIL"}  {label}{" — " + detail if detail else ""}')
    if not cond:
        failures += 1


# 1. Valid sample → output opens in python-docx and has body text
status, out = post_docx((ROOT / 'test' / 'sample_nghi_dinh.docx').read_bytes(), 'Nghị định 141.docx')
check('valid upload returns 200', status == 200, f'status={status}')
doc = Document(io.BytesIO(out))
texts = [p.text for p in doc.paragraphs if p.text.strip()]
check('output parses as docx with content', len(texts) > 5, f'{len(texts)} paragraphs')

# 2. Well-formed docx without the header table → 422 with clean message
buf = io.BytesIO()
plain = Document()
plain.add_paragraph('Tài liệu thường, không phải văn bản pháp luật.')
plain.save(buf)
status, out = post_docx(buf.getvalue(), 'plain.docx')
msg = json.loads(out)['error']
check('unrecognized structure returns 422', status == 422, f'status={status}')
check('error is clean Vietnamese text', 'Traceback' not in msg and len(msg) > 0, msg)

sys.exit(1 if failures else 0)
